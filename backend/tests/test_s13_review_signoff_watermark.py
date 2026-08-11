"""
Sprint 13 Review Notes, Sign-Off Gate, Redaction & Watermark Tests.

Verifies:
1. PII Redaction utility (owner names, survey numbers, GPS coordinates).
2. Draft watermark presence in unsigned DOCX exports, absence in signed DOCX exports.
3. Review note workflow: Analyst prepares, valuer raises note, sign refused due to open note, analyst responds, note closed, valuer signs.
4. IBBI valuer registration check on sign-off gate.
"""

from __future__ import annotations

import asyncio
import os
import uuid

import pytest
import pytest_asyncio
from docx import Document
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.models.deliverable import Deliverable
from app.models.firm import Firm
from app.models.user import User
from app.repositories import deliverableRepository, reviewNoteRepository
from app.routers import reviewNotes
from app.services.access.scope import FirmScope
from app.services.render import docxRenderer
from app.utils.redaction import redact_text

TEST_DATABASE_URL = os.environ.get("TEST_DATABASE_URL")


def test_redaction_utility():
    """Verify owner names, survey numbers, and GPS coordinates are redacted."""
    raw = "Subject property owned by Owner: Ramesh Kumar at Survey No: 42/1A located at POINT(73.8567 18.5204)."
    clean = redact_text(raw)
    assert "Ramesh Kumar" not in clean
    assert "42/1A" not in clean
    assert "73.8567" not in clean
    assert "[REDACTED_OWNER]" in clean
    assert "[REDACTED_SURVEY]" in clean
    assert "[REDACTED_COORDINATE]" in clean


def test_draft_watermark_in_docx(tmp_path):
    """Verify unsigned DOCX contains Draft watermark, signed DOCX omits it."""
    clause_plan = [{"heading": "Scope", "type": "standard_clause", "content": "Valuation scope."}]
    
    # 1. Draft render
    draft_path = docxRenderer.render(
        clause_plan=clause_plan,
        doc_type="valuation_report",
        client_name="Test Client",
        property_address="Test Address",
        computed={},
        job_id="test_draft",
        status="draft",
    )
    draft_doc = Document(draft_path)
    draft_text = "\n".join(p.text for p in draft_doc.paragraphs)
    assert "Draft — not for reliance" in draft_text

    # 2. Signed render
    signed_path = docxRenderer.render(
        clause_plan=clause_plan,
        doc_type="valuation_report",
        client_name="Test Client",
        property_address="Test Address",
        computed={},
        job_id="test_signed",
        status="signed",
    )
    signed_doc = Document(signed_path)
    signed_text = "\n".join(p.text for p in signed_doc.paragraphs)
    assert "Draft — not for reliance" not in signed_text


# ── Live DB Tests (skipped when TEST_DATABASE_URL is unconfigured) ─────────

def _rebuild_schema() -> None:
    from alembic.config import Config

    from alembic import command

    cfg = Config("alembic.ini")
    cfg.set_main_option("sqlalchemy.url", TEST_DATABASE_URL)
    command.downgrade(cfg, "base")
    command.upgrade(cfg, "head")


@pytest_asyncio.fixture
async def db():
    if not TEST_DATABASE_URL:
        pytest.skip("TEST_DATABASE_URL is not set — needs a live PostGIS (docker compose up -d)")
    await asyncio.to_thread(_rebuild_schema)
    engine = create_async_engine(TEST_DATABASE_URL)
    maker = async_sessionmaker(engine, expire_on_commit=False)
    async with maker() as session:
        yield session
    await engine.dispose()


@pytest_asyncio.fixture
async def signoff_setup(db):
    firm_id = uuid.uuid4()
    firm = Firm(id=firm_id, name="Apex Valuations", plan="pro")
    db.add(firm)
    await db.flush()

    # The scopes carry user ids, and a review note's author and assignee are
    # foreign keys to `users`. Minting bare UUIDs worked only while the table did
    # not exist — a note attributed to nobody is exactly what the sign-off trail
    # is there to prevent, so the rows are real here.
    analyst = User(firm_id=firm_id, email="analyst@apexvaluations.in", role="analyst")
    valuer = User(
        firm_id=firm_id,
        email="valuer@apexvaluations.in",
        role="valuer",
        ibbi_reg_no="IBBI/RV/02/2020/99999",
        valuer_asset_class="land_and_building",
    )
    db.add_all([analyst, valuer])
    await db.flush()

    analyst_scope = FirmScope(firm_id=firm_id, user_id=analyst.id, role="analyst")
    valuer_scope = FirmScope(
        firm_id=firm_id,
        user_id=valuer.id,
        role="valuer",
        ibbi_reg_no="IBBI/RV/02/2020/99999",
        valuer_asset_class="land_and_building",
    )

    deliv = await deliverableRepository.create_deliverable(
        db,
        scope=analyst_scope,
        doc_type="valuation_report",
        title="Signoff Gate Test Deliverable",
        sections=[{"ord": 1, "section_type": "executive_summary", "content": "Summary content"}],
    )
    return {
        "firm": firm,
        "analyst_scope": analyst_scope,
        "valuer_scope": valuer_scope,
        "deliverable": deliv,
    }


@pytest.mark.asyncio
async def test_review_note_workflow_and_signoff_gate(db, signoff_setup):
    setup = signoff_setup
    analyst: FirmScope = setup["analyst_scope"]
    valuer: FirmScope = setup["valuer_scope"]
    deliv: Deliverable = setup["deliverable"]

    # 1. Valuer raises a review note
    note = await reviewNoteRepository.create_note(
        db, scope=valuer, deliverable_id=deliv.id, note="Verify cap rate calculation"
    )
    assert note.status == "open"

    # 2. Attempt to sign deliverable is refused due to open note
    with pytest.raises(Exception) as exc_info:
        await reviewNotes.sign_deliverable(
            deliverable_id=deliv.id,
            req=reviewNotes.SignDeliverableRequest(asset_class="land_and_building"),
            db=db,
            scope=valuer,
        )
    assert "open review note" in str(exc_info.value).lower()

    # 3. Analyst responds to note
    await reviewNoteRepository.respond_note(db, scope=analyst, note_id=note.id, response_text="Cap rate verified against market coms.")
    
    # 4. Valuer closes note
    await reviewNoteRepository.close_note(db, scope=valuer, note_id=note.id)

    # 5. Now valuer signs successfully
    signed_res = await reviewNotes.sign_deliverable(
        deliverable_id=deliv.id,
        req=reviewNotes.SignDeliverableRequest(asset_class="land_and_building"),
        db=db,
        scope=valuer,
    )
    assert signed_res["status"] == "signed"
    assert signed_res["signed_by"] == str(valuer.user_id)
