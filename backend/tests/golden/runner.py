"""
Golden runner — the S1 "zero behaviour change" proof.

Runs a fixed input through the graph with the LLM replaced by a cassette, then
records an *observation*: the job type, the parse result, every computed figure,
the section sequence, and every number that reached the rendered DOCX.

`--target flat` runs the pre-S1 modules at the repository root.
`--target package` runs `backend/app/services/graph/reGraph.py`.

Identical observations across the two targets is the exit proof. Figures are
compared as exact strings, so a rupee of drift fails.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import tempfile
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path

HERE = Path(__file__).resolve().parent
BACKEND_ROOT = HERE.parents[1]
REPO_ROOT = BACKEND_ROOT.parent

CASES_DIR = HERE / "cases"
CASSETTES_DIR = HERE / "cassettes"
EXPECTED_DIR = HERE / "expected"

# Every number in the rendered document, as written.
FIGURE_RE = re.compile(r"-?\d[\d,]*(?:\.\d+)?")


def _json_safe(value):
    """
    Decimal and date render as exact strings.

    Never as floats: the whole point of S6 is that these figures are exact, and
    a golden file that stored them as doubles would compare 7499.589999999999
    against 7499.59 and either fail spuriously or, worse, pass by rounding.
    """
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    raise TypeError(f"{type(value).__name__} is not JSON serialisable")


def normalise(observation: dict) -> dict:
    """Round-trip through the encoder so observed and expected are comparable."""
    return json.loads(json.dumps(observation, default=_json_safe, ensure_ascii=False))


# ── target loading ────────────────────────────────────────────────────────────

def load_target(target: str):
    """Return (graph_module, job_types, renderer_output_dir_setter)."""
    if target == "flat":
        sys.path.insert(0, str(REPO_ROOT))
        import config  # noqa: F401
        import re_graph as graph_module
        job_types = config.ALL_JOB_TYPES
    elif target == "package":
        sys.path.insert(0, str(BACKEND_ROOT))
        from app.configs import jobTypes
        from app.services.graph import reGraph as graph_module
        job_types = jobTypes.ALL_JOB_TYPES
    else:
        raise ValueError(f"unknown target {target!r}")
    return graph_module, set(job_types)


# ── observation ───────────────────────────────────────────────────────────────

def docx_text(path: str) -> list[str]:
    from docx import Document

    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return out


def observe(final: dict, doc_path: str) -> dict:
    computed = final.get("computed") or {}
    clause_plan = final.get("clause_plan") or []
    lines = docx_text(doc_path)
    body = "\n".join(lines)

    return {
        "doc_type": final.get("doc_type"),
        "parsed_format": (final.get("parsed_data") or {}).get("format"),
        "parsed_record_count": len(((final.get("parsed_data") or {}).get("records")) or []),
        "computed": computed,
        "section_sequence": [
            {"heading": c.get("heading", ""), "type": c.get("type", "")} for c in clause_plan
        ],
        "clause_count": len(clause_plan),
        "document_figures": FIGURE_RE.findall(body),
        "document_line_count": len(lines),
        "document_text_sha256": hashlib.sha256(body.encode("utf-8")).hexdigest(),
        "doc_filename": os.path.basename(doc_path),
        "generation_errors": final.get("generation_errors"),
    }


# ── run ───────────────────────────────────────────────────────────────────────

def run_case(case_name: str, target: str, mode: str, output_dir: Path) -> dict:
    from cassette import Cassette  # local to this directory

    case = json.loads((CASES_DIR / f"{case_name}.json").read_text())
    output_dir.mkdir(parents=True, exist_ok=True)

    # The renderer reads OUTPUT_DIR at import time.
    os.environ["OUTPUT_DIR"] = str(output_dir)
    os.environ.setdefault("GROQ_API_KEY", "golden-run-no-live-calls")
    os.environ.setdefault("JWT_SECRET", "golden-run-no-tokens-are-minted-here")
    os.environ.pop("S3_BUCKET", None)      # keep upload on the file:// branch
    os.environ.pop("GEMINI_API_KEY", None)  # keep vision on the no-letterhead branch

    graph_module, job_types = load_target(target)
    tape = Cassette(CASSETTES_DIR / f"{case_name}.json", mode=mode, job_types=job_types)
    graph_module._chat = tape.chat

    final = graph_module.app.invoke({
        # S8: the evidence the fixture's property carries. In production this is
        # assembled by generationService from a scoped repository call; here the
        # fixture supplies it, and the gate evaluates it exactly the same way.
        "property_id": case.get("property_id"),
        "evidence_bundle": case.get("evidence"),
        "evidence_checked": False,
        "evidence_missing": None,
        "_blocked": False,
        "_scope": None,
        "raw_instructions": case["instructions"],
        "raw_property_data": case.get("property_data", ""),
        "job_type": case.get("job_type"),
        "_job_id": case["job_id"],
        "doc_type": None, "client_name": None,
        "property_address": None, "property_type": None,
        "purpose": None, "special_notes": None,
        "parsed_data": None, "computed": None,
        "re_research": None, "header_image_path": None,
        "structure_plan": None, "structure_attempt": 0,
        "critic_feedback": None, "_critic_approved": False,
        "section_index": 0, "drafted_sections": None,
        "clause_plan": None,
        "render_attempt": 0, "render_error": None,
        "doc_path": None, "doc_url": None,
        "generation_errors": None,
    })
    tape.save()

    doc_path = final.get("doc_path")
    if not doc_path:
        raise RuntimeError(f"{case_name}: nothing rendered — {final.get('render_error')}")
    obs = observe(final, doc_path)
    obs["llm_call_count"] = len(tape.calls)
    return normalise(obs)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--target", choices=("flat", "package"), required=True)
    ap.add_argument("--mode", choices=("replay", "synthesise", "record"), default="replay")
    ap.add_argument("--write-expected", action="store_true")
    ap.add_argument("--out", default=None, help="where to write observations")
    ap.add_argument("cases", nargs="*", default=None)
    args = ap.parse_args()

    names = args.cases or sorted(p.stem for p in CASES_DIR.glob("*.json"))
    out_dir = Path(args.out) if args.out else Path(tempfile.mkdtemp(prefix="golden-"))

    failures = 0
    for name in names:
        obs = run_case(name, args.target, args.mode, out_dir / name)
        target_path = (EXPECTED_DIR if args.write_expected else out_dir) / f"{name}.json"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_text(json.dumps(obs, indent=2, ensure_ascii=False, default=_json_safe) + "\n")

        if args.write_expected:
            print(f"recorded  {name}  ({obs['clause_count']} sections, "
                  f"{len(obs['document_figures'])} figures)")
        else:
            expected = json.loads((EXPECTED_DIR / f"{name}.json").read_text())
            if expected == obs:
                print(f"ok        {name}")
            else:
                failures += 1
                print(f"MISMATCH  {name}")
                for key in sorted(set(expected) | set(obs)):
                    if expected.get(key) != obs.get(key):
                        print(f"            {key}: expected {expected.get(key)!r}")
                        print(f"            {key}: actual   {obs.get(key)!r}")
    print(f"\nobservations in {out_dir}")
    return 1 if failures else 0


def _pin_hash_seed() -> None:
    """
    `intake_node` builds its prompt with `', '.join(ALL_JOB_TYPES)` over a set, so
    the prompt text — and therefore the whole run — changes between processes under
    Python's string hash randomisation. Pinning the seed makes a golden run
    reproducible without touching the graph. The underlying instability is a real
    defect (it also defeats prompt caching) and belongs to S10, where prompts
    become data.
    """
    if os.environ.get("PYTHONHASHSEED") != "0":
        os.environ["PYTHONHASHSEED"] = "0"
        os.execv(sys.executable, [sys.executable, str(Path(__file__).resolve()), *sys.argv[1:]])


if __name__ == "__main__":
    _pin_hash_seed()
    sys.path.insert(0, str(HERE))
    raise SystemExit(main())
