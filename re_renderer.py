"""
RE Renderer — Real Estate Factory
Deterministic python-docx renderer. No LLM calls.
Handles: executive_summary, property_description, market_analysis,
         valuation_approach, due_diligence_check, construction_stage,
         summary_table, unit_table, recommendations, conclusion, standard_clause.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config

OUTPUT_DIR = os.environ.get("OUTPUT_DIR", "output")

NAVY        = RGBColor(0x1B, 0x2A, 0x4A)
GOLD        = RGBColor(0xB8, 0x96, 0x0C)
DARK_GREY   = RGBColor(0x3A, 0x3A, 0x3A)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)


def _shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _heading(doc, text, level=1, colour=NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = colour
        run.font.bold = True
    return h


def _para(doc, text, bold=False, size=11, colour=DARK_GREY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    return p


def _render_cover(doc, doc_type, client_name, property_address):
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(doc_type.replace("_", " ").upper())
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    if client_name:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = s.add_run(f"Prepared for: {client_name}")
        r2.font.size = Pt(13)
        r2.font.color.rgb = DARK_GREY

    if property_address:
        s2 = doc.add_paragraph()
        s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = s2.add_run(f"Property: {property_address}")
        r3.font.size = Pt(11)
        r3.font.color.rgb = DARK_GREY
        r3.italic = True

    doc.add_page_break()


def _render_summary_table(doc, clause):
    _heading(doc, clause.get("heading", "Summary"), level=2)
    content = clause.get("content", "")
    # Parse key: value lines into a 2-col table
    lines = [l for l in content.split("\n") if l.strip()]
    kv = []
    for line in lines:
        if "|" in line:
            # Already table-formatted — render as text
            _para(doc, line)
            continue
        if ":" in line:
            parts = line.split(":", 1)
            kv.append((parts[0].strip(), parts[1].strip()))
        else:
            _para(doc, line)

    if kv:
        table = doc.add_table(rows=len(kv), cols=2)
        table.style = "Table Grid"
        for i, (k, v) in enumerate(kv):
            table.cell(i, 0).text = k
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            _shade_cell(table.cell(i, 0), "E8F0FE")
            table.cell(i, 1).text = v
    doc.add_paragraph()


def _render_unit_table(doc, clause):
    _heading(doc, clause.get("heading", "Details"), level=2)
    content = clause.get("content", "")
    lines   = [l for l in content.split("\n") if l.strip()]
    if not lines:
        return

    rows_data = [l.split("|") for l in lines]
    if not rows_data:
        return

    cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows_data):
        for j, cell_val in enumerate(row):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = cell_val.strip()
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    _shade_cell(cell, "1B2A4A")
                    cell.paragraphs[0].runs[0].font.color.rgb = WHITE
    doc.add_paragraph()


def _render_standard(doc, clause):
    heading = clause.get("heading", "")
    content = clause.get("content", "")
    if heading:
        _heading(doc, heading, level=2)
    if content:
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0].isdigit() and ". " in line[:4]:
                doc.add_paragraph(line, style="List Number")
            else:
                _para(doc, line)
    doc.add_paragraph()


_RENDERERS = {
    "summary_table":        _render_summary_table,
    "unit_table":           _render_unit_table,
    "executive_summary":    _render_standard,
    "property_description": _render_standard,
    "market_analysis":      _render_standard,
    "valuation_approach":   _render_standard,
    "due_diligence_check":  _render_standard,
    "construction_stage":   _render_standard,
    "recommendations":      _render_standard,
    "conclusion":           _render_standard,
    "standard_clause":      _render_standard,
    "numbered_clause":      _render_standard,
}


def render(
    clause_plan: list,
    doc_type: str,
    client_name: str,
    property_address: str,
    computed: dict,
    header_image_path: str = None,
    job_id: str = "job",
) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    if header_image_path and os.path.exists(header_image_path):
        doc.add_picture(header_image_path, width=Inches(6.0))
        doc.add_paragraph()

    _render_cover(doc, doc_type, client_name, property_address)

    for clause in (clause_plan or []):
        clause_type = clause.get("type", "standard_clause")
        fn = _RENDERERS.get(clause_type, _render_standard)
        fn(doc, clause)

    out_path = os.path.join(OUTPUT_DIR, f"{job_id}_{doc_type}.docx")
    doc.save(out_path)
    return out_path
